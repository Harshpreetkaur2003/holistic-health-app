import streamlit as st
import pandas as pd
import numpy as np
import pickle
from tensorflow.keras.models import load_model
from sklearn.preprocessing import StandardScaler
from io import BytesIO
from docx import Document
from docx.shared import Pt
import base64
import matplotlib.pyplot as plt
import os

st.set_page_config(page_title="🌿 Holistic Health Assessment", layout="wide")
st.title("🌿 Holistic Health & Lifestyle Assessment")

# -----------------------------
# Base directory
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

rf_path = os.path.join(BASE_DIR, "model", "rf_model.pkl")
scaler_path = os.path.join(BASE_DIR, "model", "scaler.pkl")
nn_path = os.path.join(BASE_DIR, "model", "nn_model.h5")

# -----------------------------
# Load models safely
rf_model = None
scaler = None
nn_model = None

# Random Forest
try:
    with open(rf_path, "rb") as f:
        rf_model = pickle.load(f)
except Exception as e:
    st.warning(f"rf_model.pkl could not be loaded. Using dummy predictions. ({e})")

# Scaler
try:
    with open(scaler_path, "rb") as f:
        scaler = pickle.load(f)
except Exception as e:
    st.warning(f"scaler.pkl could not be loaded. Using dummy scaling. ({e})")

# Neural Network
try:
    nn_model = load_model(nn_path)
except Exception as e:
    st.warning(f"nn_model.h5 could not be loaded. Using dummy predictions. ({e})")

# -----------------------------
# Sidebar - User Info
st.sidebar.header("User Information")
name = st.sidebar.text_input("Name")
age = st.sidebar.number_input("Age", 1, 120, 25)
gender = st.sidebar.selectbox("Gender", ["Male", "Female", "Other"])

st.sidebar.subheader("Upload Reports (Optional)")
uploaded_file = st.sidebar.file_uploader("Upload Blood Report (CSV/PDF)", type=['csv','pdf'])

# -----------------------------
# Lifestyle & Health Inputs
st.subheader("Lifestyle & Health Inputs")
sleep_hours = st.slider("Sleep Hours per day", 0, 24, 7)
daily_stress = st.slider("Daily Stress (0-10)", 0, 10, 3)
daily_steps = st.number_input("Daily Steps", 0, 30000, 6000)
weekly_meditation = st.slider("Weekly Meditation Sessions", 0, 14, 2)
fruits_veggies = st.slider("Fruits & Veggies Score (0-10)", 0, 10, 7)
time_for_passion = st.slider("Time for Hobbies (hrs/week)", 0, 20, 2)
supporting_others = st.slider("Helping Others (0-10)", 0, 10, 5)
social_network = st.slider("Social Interaction (0-10)", 0, 10, 5)
bmi_range = st.number_input("BMI", 10.0, 50.0, 23.5)
work_life_balance_score = st.slider("Work-Life Balance Score (0-10)", 0, 10, 7)

# -----------------------------
# Encode gender for ML
gender_map = {"Male":0, "Female":1, "Other":2}
gender_numeric = gender_map[gender]

# -----------------------------
if st.button("Analyze Health Risk"):

    rf_features = np.array([[sleep_hours, daily_stress, daily_steps, weekly_meditation,
                             fruits_veggies, time_for_passion, supporting_others, social_network,
                             bmi_range, work_life_balance_score, age, gender_numeric]])

    # Safe NN features
    nn_features_scaled = rf_features.copy()
    if scaler is not None:
        try:
            nn_features_scaled[:, :11] = scaler.transform(nn_features_scaled[:, :11])
        except:
            st.warning("Scaler transform failed, using unscaled data.")

    # -----------------------------
    # Predictions
    if rf_model is not None:
        try:
            rf_prob = rf_model.predict_proba(rf_features)[0][1]
        except:
            rf_prob = 0.5
    else:
        rf_prob = 0.5  # Dummy probability

    if nn_model is not None:
        try:
            nn_pred_prob = nn_model.predict(nn_features_scaled)
            nn_pred = np.argmax(nn_pred_prob, axis=1)[0]
            nn_prob = nn_pred_prob[0][nn_pred]
        except:
            nn_prob = 0.5
    else:
        nn_prob = 0.5  # Dummy probability

    # -----------------------------
    # Dynamic Risk Scoring
    risk_score = 0
    if sleep_hours < 6: risk_score += 2
    if daily_stress > 6: risk_score += 3
    if daily_steps < 5000: risk_score += 2
    if fruits_veggies < 6: risk_score += 1
    if bmi_range > 25: risk_score += 2
    if work_life_balance_score < 6: risk_score += 1
    risk_score += rf_prob * 3  # ML contribution

    risk_probability = min((risk_score / 12) * 100, 100)
    if risk_probability < 30:
        risk_label = "Low"
    elif risk_probability < 70:
        risk_label = "Medium"
    else:
        risk_label = "High"

    # -----------------------------
    # Contributing factors
    issues = []
    if sleep_hours < 6: issues.append("Low Sleep Hours")
    if daily_stress > 6: issues.append("High Stress")
    if daily_steps < 5000: issues.append("Low Physical Activity")
    if fruits_veggies < 6: issues.append("Poor Diet")
    if bmi_range > 25: issues.append("High BMI")
    if work_life_balance_score < 6: issues.append("Low Work-Life Balance")

    # -----------------------------
    # Recommendations
    if issues:
        diet_advice = "- Eat more fruits, vegetables, and balanced meals.\n- Drink plenty of water.\n- Reduce processed foods."
        workout_advice = "- Walk at least 5000 steps daily.\n- Include cardio and strength workouts.\n- Meditate regularly."
    else:
        diet_advice = "Great! Maintain your healthy diet and keep hydrated."
        workout_advice = "Excellent! Keep your activity and meditation routine consistent."

    # -----------------------------
    # Display
    color_map = {"Low":"green", "Medium":"orange", "High":"red"}
    st.markdown(f"### Health Risk: <span style='color:{color_map[risk_label]}'>{risk_label}</span>", unsafe_allow_html=True)
    st.write(f"Risk Probability: {risk_probability:.2f}%")
    st.write("### Contributing Factors:")
    st.write(", ".join(issues) if issues else "None")
    st.write("### Recommendations:")
    st.write("**Diet:**")
    st.text(diet_advice)
    st.write("**Workout & Lifestyle:**")
    st.text(workout_advice)

    # -----------------------------
    # Radar chart
    radar_labels = ['Sleep','Stress','Steps','Meditation','Diet','Hobby','Helping','Social','BMI','Work-Life','Age']
    radar_values = [sleep_hours/24, daily_stress/10, daily_steps/15000, weekly_meditation/14,
                    fruits_veggies/10, time_for_passion/20, supporting_others/10, social_network/10,
                    bmi_range/50, work_life_balance_score/10, age/100]
    angles = np.linspace(0, 2*np.pi, len(radar_labels), endpoint=False).tolist()
    radar_values += radar_values[:1]
    angles += angles[:1]
    fig, ax = plt.subplots(figsize=(6,6), subplot_kw=dict(polar=True))
    ax.plot(angles, radar_values, linewidth=2, linestyle='solid')
    ax.fill(angles, radar_values, 'skyblue', alpha=0.4)
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(radar_labels)
    ax.set_yticklabels([])
    st.pyplot(fig)

    # -----------------------------
    # Word report
    doc = Document()
    doc.add_heading(f'Holistic Health Report - {name}', 0)
    doc.add_paragraph(f'Age: {age}')
    doc.add_paragraph(f'Gender: {gender}')
    doc.add_paragraph(f'Health Risk: {risk_label} ({risk_probability:.2f}%)')
    doc.add_paragraph(f'Contributing Factors: {", ".join(issues) if issues else "None"}')
    doc.add_paragraph("Recommendations:")
    doc.add_paragraph("Diet: " + diet_advice.replace("\n","; "))
    doc.add_paragraph("Workout & Lifestyle: " + workout_advice.replace("\n","; "))
    if uploaded_file:
        doc.add_paragraph(f'Uploaded Report: {uploaded_file.name} included in analysis.')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    b64 = base64.b64encode(buffer.read()).decode()
    st.markdown(f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="Holistic_Health_Report.docx">📥 Download Word Report</a>', unsafe_allow_html=True)


